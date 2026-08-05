"""绿色碳排 Word 报告生成。

借鉴 `_ext_src/building-energy-esg-fm` 的 `v2_report_export.py` 架构：
上下文组装（这里直接读取 `CarbonAnalysis.result_json` 与顶层列）→ python-docx 生成 →
缺依赖降级为纯文本。纯生成模块，不触碰数据库。
"""

from __future__ import annotations

from io import BytesIO


def _fmt(value, digits: int = 4) -> str:
    """数字格式化：去掉无意义的末尾 0，None/非法值回退为 —。"""
    if value is None or value == "":
        return "—"
    try:
        return f"{float(value):.{digits}f}".rstrip("0").rstrip(".")
    except (TypeError, ValueError):
        return str(value)


def _pct(share, digits: int = 1) -> str:
    if share is None:
        return "—"
    try:
        return f"{float(share) * 100:.{digits}f}%"
    except (TypeError, ValueError):
        return "—"


def _fmt_datetime(dt) -> str:
    if dt is None:
        return "—"
    if hasattr(dt, "astimezone"):
        try:
            return dt.astimezone().strftime("%Y-%m-%d %H:%M")
        except Exception:
            pass
    return str(dt)[:16]


def _esg_summary_lines(result: dict, analysis) -> tuple[str, str, str]:
    """E/S/G 三行摘要，全部使用本项目真实字段（不编造规范条款号）。"""
    stages = result.get("stages", [])
    suggestions = result.get("suggestions", [])
    has_unverified = bool(result.get("has_unverified_factors", False))
    total = analysis.total_emission or 0.0
    area = analysis.area_m2
    intensity = round(total / area, 4) if area and area > 0 else None

    stage_parts = (
        " / ".join(f"{s.get('stage')} {s.get('stage_name')} {s.get('share', 0) * 100:.0f}%" for s in stages)
        or "无分阶段数据"
    )
    e = f"总排放 {_fmt(total)} tCO2e" + (
        f"、面积强度 {_fmt(intensity)} tCO2e/m²" if intensity is not None else ""
    ) + f"；分阶段占比 {stage_parts}。"
    s = (
        f"共 {len(suggestions)} 条减排建议；绿色施工管理参照 GB/T 50905-2014 绿色施工方向落实，"
        "具体以项目绿色施工专项方案为准。"
    )
    g = (
        f"因子核证状态：{'部分因子待核证' if has_unverified else '已核证'}；"
        f"核算模式：{'演示因子' if analysis.is_simulated else '核证因子'}；"
        f"因子库版本 {analysis.factor_version or 'unavailable'}。"
    )
    return e, s, g


def _fallback_txt(analysis, project_name: str) -> str:
    preview = analysis.report_preview
    if preview:
        return f"《建筑碳排放核算报告（GB/T 51366-2019）》\n项目：{project_name}（{analysis.project_id}）\n\n{preview}"
    return f"碳排报告 {analysis.id}（python-docx 未安装，无法生成 Word）"


def build_report_docx(analysis, project_name: str) -> tuple[bytes, str, str]:
    """生成《建筑碳排放核算报告》Word 文档。

    返回 (body_bytes, filename, media_type)；python-docx 缺失时降级为 .txt。
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
    except ImportError:
        body = _fallback_txt(analysis, project_name).encode("utf-8")
        return body, f"carbon_report_{analysis.id}.txt", "text/plain; charset=utf-8"

    result = analysis.result_json if isinstance(analysis.result_json, dict) else {}
    stages = result.get("stages", [])
    items = result.get("items", [])
    contributors = result.get("top_contributors", [])
    suggestions = result.get("suggestions", [])
    warnings = result.get("factor_warnings", [])
    has_unverified = bool(result.get("has_unverified_factors", False))

    total = analysis.total_emission or 0.0
    area = analysis.area_m2
    intensity = round(total / area, 4) if area and area > 0 else None

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")

    def para(text: str, *, size: float = 10.5, bold: bool = False, align=None, color=None):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = RGBColor(*color)
        return p

    def heading(text: str):
        return para(text, size=12, bold=True, color=(0x1F, 0x3B, 0x73))

    def add_table(headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.text = ""
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = "" if value is None else str(value)
        return table

    # 页眉页脚
    header = doc.sections[0].header.paragraphs[0]
    header.text = "筑智共生 BuildWise · 绿色建造碳排核算"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "本报告由因子法自动生成，仅供绿色建造管理参考。"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)

    # 标题与元信息
    para("建筑碳排放核算报告（GB/T 51366-2019）", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(
        f"项目：{project_name}（{analysis.project_id}）　|　生成时间：{_fmt_datetime(analysis.created_at)}　|　分析编号：{analysis.id}",
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph()

    # 一、报告说明
    heading("一、报告说明")
    para("本报告依据 GB/T 51366-2019《建筑碳排放计算标准》，采用因子法核算施工阶段碳排放：排放 = Σ(活动数据 × 排放因子)。")
    para("核算范围覆盖 A1-A3（建材生产）、A4（建材运输）、A5（施工过程）三个阶段，数据来源为项目录入的活动数据与因子库排放因子。")
    if analysis.is_simulated:
        para(
            "本报告使用未完全核证的演示因子数据，结果仅供演示与决策参考，正式核算需替换为经核证的因子数据。",
            color=(0xC2, 0x41, 0x0C),
        )

    # 二、核算结果总览
    heading("二、核算结果总览")
    overview = [
        ("分析编号", analysis.id),
        ("项目名称", project_name),
        ("生成时间", _fmt_datetime(analysis.created_at)),
        ("建筑面积（m²）", _fmt(area) if area else "未填写"),
        ("核算范围", analysis.scope or "施工阶段 A1-A3 / A4 / A5（GB/T 51366-2019 因子法）"),
        ("总排放（tCO2e）", _fmt(total)),
        ("单位建筑面积排放（tCO2e/m²）", _fmt(intensity) if intensity is not None else "未填写"),
        ("因子库版本", analysis.factor_version or "unavailable"),
        ("核算模式", "演示因子（待核证）" if analysis.is_simulated else "核证因子"),
    ]
    add_table(["项目", "数值"], overview)
    doc.add_paragraph()

    # 三、分阶段排放
    heading("三、分阶段排放")
    stage_rows = [
        [
            s.get("stage"),
            s.get("stage_name"),
            _fmt(s.get("emission")),
            _pct(s.get("share")),
            s.get("items_count", 0),
        ]
        for s in stages
    ]
    add_table(["阶段", "名称", "排放（tCO2e）", "占比", "记录数"], stage_rows)
    doc.add_paragraph()

    # 四、主要贡献项
    heading("四、主要贡献项")
    item_source = {(i.get("stage"), i.get("code")): i.get("factor_source", "") for i in items}
    contributor_rows = [
        [
            c.get("name"),
            c.get("stage"),
            _fmt(c.get("emission")),
            _pct(c.get("share")),
            item_source.get((c.get("stage"), c.get("code")), "") or "—",
        ]
        for c in contributors
    ]
    if contributor_rows:
        add_table(["项目", "阶段", "排放（tCO2e）", "占比", "因子来源"], contributor_rows)
    else:
        para("暂无贡献项数据。")
    doc.add_paragraph()

    # 五、减排建议
    heading("五、减排建议")
    if suggestions:
        for index, suggestion in enumerate(suggestions, start=1):
            para(f"{index}. {suggestion}")
    else:
        para("暂无减排建议。")
    doc.add_paragraph()

    # 六、E/S/G 摘要与核证说明
    heading("六、E/S/G 摘要")
    e, s, g = _esg_summary_lines(result, analysis)
    for label, text in (("E（环境）", e), ("S（社会）", s), ("G（治理）", g)):
        p = doc.add_paragraph()
        run = p.add_run(f"{label}：")
        run.bold = True
        p.add_run(text)
    doc.add_paragraph()

    if warnings or has_unverified:
        heading("核证说明")
        if has_unverified:
            para("结果使用了未核证的演示排放因子，仅供演示；正式核算需替换为经核证的因子数据。", color=(0xC2, 0x41, 0x0C))
        for warning in warnings:
            para(f"· {warning}")

    buffer = BytesIO()
    doc.save(buffer)
    return (
        buffer.getvalue(),
        f"carbon_report_{analysis.id}.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
