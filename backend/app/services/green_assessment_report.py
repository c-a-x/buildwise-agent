"""四节一环保评估 Word 报告生成。

沿用 `carbon_report.py` 架构：读取 `GreenAssessment.result_json` 与顶层列 → python-docx 生成 →
缺依赖降级为纯文本。纯生成模块，不触碰数据库。
"""

from __future__ import annotations

from io import BytesIO


def _fmt(value, digits: int = 1) -> str:
    if value is None or value == "":
        return "—"
    try:
        return f"{float(value):.{digits}f}".rstrip("0").rstrip(".")
    except (TypeError, ValueError):
        return str(value)


def _fmt_datetime(dt) -> str:
    if dt is None:
        return "—"
    if hasattr(dt, "astimezone"):
        try:
            return dt.astimezone().strftime("%Y-%m-%d %H:%M")
        except Exception:
            pass
    return str(dt)[:16]


def _fallback_txt(assessment, project_name: str) -> str:
    preview = assessment.report_preview
    if preview:
        return f"《四节一环保评估报告》\n项目：{project_name}（{assessment.project_id}）\n\n{preview}"
    return f"四节一环保评估 {assessment.id}（python-docx 未安装，无法生成 Word）"


def build_assessment_docx(assessment, project_name: str) -> tuple[bytes, str, str]:
    """生成《四节一环保评估报告》Word 文档。

    返回 (body_bytes, filename, media_type)；python-docx 缺失时降级为 .txt。
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
    except ImportError:
        body = _fallback_txt(assessment, project_name).encode("utf-8")
        return body, f"green_assessment_{assessment.id}.txt", "text/plain; charset=utf-8"

    result = assessment.result_json if isinstance(assessment.result_json, dict) else {}
    dimensions = result.get("dimensions", [])
    total = assessment.total_score or 0.0
    level = assessment.level or "—"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")

    def para(text: str, *, size: float = 10.5, bold: bool = False, align=None, color=None):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = RGBColor(*color)
        return p

    def heading(text: str):
        return para(text, size=12, bold=True, color=(0x1F, 0x3B, 0x73))

    def add_table(headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.text = ""
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = "" if value is None else str(value)
        return table

    header = doc.sections[0].header.paragraphs[0]
    header.text = "筑智共生 BuildWise · 绿色建造四节一环保评估"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "本报告由评估算法自动生成，仅供绿色建造管理参考。"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)

    para("四节一环保评估报告", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(
        f"项目：{project_name}（{assessment.project_id}）　|　生成时间：{_fmt_datetime(assessment.created_at)}　|　评估编号：{assessment.id}",
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph()

    heading("一、评估说明")
    para("依据《绿色施工导则》与四节一环保（节材、节水、节能、节地、环境保护）要求，对项目绿色施工水平进行量化评估。")
    para("每维度由若干指标构成，各指标按目标达成度折算 0~100 分，维度均分后按 20% 权重汇总总分。")
    if assessment.is_simulated:
        para("本评估存在未填写的指标（按 0 分计），结果仅供演示，请补充完整数据后重新评估。", color=(0xC2, 0x41, 0x0C))

    heading("二、评估结果总览")
    overview = [
        ("评估编号", assessment.id),
        ("项目名称", project_name),
        ("生成时间", _fmt_datetime(assessment.created_at)),
        ("建筑面积（m²）", _fmt(assessment.area_m2) if assessment.area_m2 else "未填写"),
        ("总分", _fmt(total)),
        ("等级", level),
    ]
    add_table(["项目", "数值"], overview)
    doc.add_paragraph()

    heading("三、分维度得分")
    dimension_rows = []
    for item in dimensions:
        detail = "；".join(
            f"{m.get('name')} {_fmt(m.get('score'))}"
            for m in item.get("metrics", [])
        )
        dimension_rows.append([item.get("name"), _fmt(item.get("score")), detail])
    if dimension_rows:
        add_table(["维度", "得分", "指标明细"], dimension_rows)
    else:
        para("暂无维度数据。")
    doc.add_paragraph()

    heading("四、提升方向")
    if dimensions:
        low = sorted(dimensions, key=lambda item: float(item.get("score") or 0))[:2]
        for item in low:
            para(f"· {item.get('name')} 得分偏低（{_fmt(item.get('score'))}），建议优先制定专项提升措施。")
    para("绿色施工具体措施以项目绿色施工专项方案为准，相关标准参照《绿色施工导则》执行。")

    buffer = BytesIO()
    doc.save(buffer)
    return (
        buffer.getvalue(),
        f"green_assessment_{assessment.id}.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
