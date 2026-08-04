from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from app.knowledge.parsers import KnowledgeParseError, parse_knowledge_file


def _write_minimal_pdf(path: Path, text_lines: list[str]) -> None:
    stream = "BT /F1 12 Tf 72 720 Td " + " ".join(
        f"({line}) Tj 0 -20 Td" for line in text_lines
    ) + " ET"
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        f"<< /Length {len(stream.encode('ascii'))} >>\nstream\n{stream}\nendstream".encode("ascii"),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    pdf = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = [0]
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{number} 0 obj\n".encode("ascii"))
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")
    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    pdf.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode(
            "ascii"
        )
    )
    path.write_bytes(pdf)


def test_json_import_preserves_clause_metadata(tmp_path: Path) -> None:
    source_path = tmp_path / "authorized-safety.json"
    source_path.write_text(
        json.dumps(
            {
                "document_id": "DOC-SAFETY-2026",
                "title": "施工现场安全制度",
                "source": "已授权项目制度",
                "version": "2026.1",
                "effective_date": "2026-01-01",
                "category": "个人防护",
                "clauses": [
                    {
                        "article": "第12条",
                        "content": "进入施工现场的人员应正确佩戴安全帽。",
                        "metadata": {"hazard_types": ["no_helmet"], "keywords": ["安全帽"]},
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    clauses = parse_knowledge_file(source_path)

    assert len(clauses) == 1
    clause = clauses[0]
    assert clause.document_id == "DOC-SAFETY-2026"
    assert clause.source == "已授权项目制度"
    assert clause.title == "施工现场安全制度"
    assert clause.article == "第12条"
    assert clause.category == "个人防护"
    assert clause.content == "进入施工现场的人员应正确佩戴安全帽。"
    assert clause.version == "2026.1"
    assert clause.effective_date == "2026-01-01"
    assert clause.metadata == {"hazard_types": ["no_helmet"], "keywords": ["安全帽"]}


def test_docx_import_splits_only_explicit_articles(tmp_path: Path) -> None:
    source_path = tmp_path / "authorized-safety.docx"
    document = Document()
    document.add_paragraph("第12条")
    document.add_paragraph("进入施工现场的人员应正确佩戴安全帽。")
    document.add_paragraph("第13条")
    document.add_paragraph("临边作业面应设置防护栏杆。")
    document.save(source_path)

    clauses = parse_knowledge_file(
        source_path,
        source="已授权规范",
        title="施工安全规范",
        category="施工安全",
        version="2026",
        effective_date="2026-02-01",
    )

    assert [(clause.article, clause.content) for clause in clauses] == [
        ("第12条", "进入施工现场的人员应正确佩戴安全帽。"),
        ("第13条", "临边作业面应设置防护栏杆。"),
    ]
    assert all(clause.source == "已授权规范" for clause in clauses)
    assert all(clause.title == "施工安全规范" for clause in clauses)


def test_pdf_import_uses_article_heading_and_preserves_source(tmp_path: Path) -> None:
    source_path = tmp_path / "authorized-safety.pdf"
    _write_minimal_pdf(source_path, ["Article 12", "Wear hard hats on site."])

    clauses = parse_knowledge_file(
        source_path,
        source="Authorized PDF",
        title="Safety Rules",
        category="PPE",
        version="2026",
    )

    assert len(clauses) == 1
    assert clauses[0].article == "Article 12"
    assert clauses[0].content == "Wear hard hats on site."
    assert clauses[0].source == "Authorized PDF"


def test_articleless_documents_are_rejected_without_fabricating_article(tmp_path: Path) -> None:
    source_path = tmp_path / "articleless.docx"
    document = Document()
    document.add_paragraph("This text has no article heading.")
    document.save(source_path)

    with pytest.raises(KnowledgeParseError, match="article"):
        parse_knowledge_file(source_path, source="Authorized", title="Rules", category="Safety")


def test_unsupported_extension_is_rejected(tmp_path: Path) -> None:
    source_path = tmp_path / "notes.txt"
    source_path.write_text("第1条\n内容", encoding="utf-8")

    with pytest.raises(KnowledgeParseError, match="Unsupported"):
        parse_knowledge_file(source_path)
